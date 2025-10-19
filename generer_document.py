#!/usr/bin/env python3
"""
Générateur de documents Word générique
Peut être utilisé pour n'importe quel thème
"""

import json
import os
import argparse
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Dictionnaire des traductions des titres par thème
THEME_TRANSLATIONS = {
    "animaux": {
        "francais": "Animaux",
        "anglais": "Animals"
    },
    "corps_humain": {
        "francais": "Parties du corps",
        "anglais": "Body parts"
    },
    "meteo": {
        "francais": "Météo",
        "anglais": "Weather"
    },
    "salon": {
        "francais": "Salon",
        "anglais": "Living room"
    },
    "toilette": {
        "francais": "Salle de bain",
        "anglais": "Bathroom"
    }
}

def load_selection_config(theme_name: str) -> dict:
    """Charge la configuration de sélection d'un thème"""
    config_path = f"themes/{theme_name}/selection.json"

    if not os.path.exists(config_path):
        raise FileNotFoundError(f"Configuration de sélection non trouvée: {config_path}")

    with open(config_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def create_word_document(theme_name: str, langue: str = "francais") -> str:
    """
    Crée un document Word pour un thème donné

    Args:
        theme_name: Nom du thème
        langue: Langue du document ("francais" ou "anglais")

    Returns:
        Chemin du document créé
    """
    # Charger la configuration
    config = load_selection_config(theme_name)

    # Obtenir le titre traduit
    if theme_name in THEME_TRANSLATIONS:
        titre_traduit = THEME_TRANSLATIONS[theme_name][langue]
    else:
        titre_traduit = config['titre']  # Fallback sur le titre original

    print("=" * 80)
    print(f"📄 GÉNÉRATION DU DOCUMENT WORD ({langue.upper()})")
    print("=" * 80)

    # Créer le document
    doc = Document()

    # Titre principal
    title = doc.add_heading(titre_traduit, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Espacement
    doc.add_paragraph()

    # Créer le tableau
    elements = config['elements']
    colonnes = config.get('colonnes', 3)
    lignes = (len(elements) + colonnes - 1) // colonnes  # Arrondi vers le haut

    table = doc.add_table(rows=lignes, cols=colonnes)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remplir le tableau
    photos_dir = f"themes/{theme_name}/photos"

    for i, element in enumerate(elements):
        row = i // colonnes
        col = i % colonnes
        cell = table.cell(row, col)

        # Ajouter l'image
        image_path = os.path.join(photos_dir, element['image_selectionnee'])

        if os.path.exists(image_path):
            try:
                # Redimensionner l'image pour qu'elle rentre dans la cellule
                cell_paragraph = cell.paragraphs[0]
                run = cell_paragraph.runs[0] if cell_paragraph.runs else cell_paragraph.add_run()

                # Ajouter l'image
                picture = run.add_picture(image_path, width=Inches(1.5))

                # Centrer l'image
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Ajouter les textes selon la langue
                cell_paragraph = cell.add_paragraph()
                cell_paragraph.add_run(element['nom_macedonien']).bold = True
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                cell_paragraph = cell.add_paragraph()
                if langue == "anglais" and 'nom_anglais' in element:
                    cell_paragraph.add_run(f"({element['nom_anglais']})").italic = True
                    print(f"   ✅ {element['nom_anglais']}: {element['image_selectionnee']}")
                else:
                    cell_paragraph.add_run(f"({element['nom_francais']})").italic = True
                    print(f"   ✅ {element['nom_francais']}: {element['image_selectionnee']}")
                cell_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            except Exception as e:
                nom_affiche = element['nom_anglais'] if langue == "anglais" and 'nom_anglais' in element else element['nom_francais']
                print(f"   ❌ Erreur image {nom_affiche}: {e}")
                nom_fallback = element['nom_anglais'] if langue == "anglais" and 'nom_anglais' in element else element['nom_francais']
                cell.text = f"{element['nom_macedonien']}\n({nom_fallback})"
        else:
            print(f"   ⚠️  Image manquante: {image_path}")
            nom_fallback = element['nom_anglais'] if langue == "anglais" and 'nom_anglais' in element else element['nom_francais']
            cell.text = f"{element['nom_macedonien']}\n({nom_fallback})"

    # Sauvegarder le document
    output_path = f"themes/{theme_name}/{titre_traduit}.docx"
    doc.save(output_path)

    print(f"\n✅ Document créé avec succès !")
    print(f"📄 Titre : {titre_traduit}")
    print(f"📊 {len(elements)} éléments")
    print(f"📐 Tableau {lignes} × {colonnes}")
    print(f"📁 Emplacement : {os.path.abspath(output_path)}")

    return output_path

def create_bilingual_documents(theme_name: str) -> list:
    """
    Crée les documents Word en français et en anglais pour un thème donné

    Args:
        theme_name: Nom du thème

    Returns:
        Liste des chemins des documents créés
    """
    documents_crees = []

    print(f"🌍 Génération des documents bilingues pour le thème '{theme_name}'")
    print("=" * 80)

    # Générer le document français
    try:
        doc_fr = create_word_document(theme_name, "francais")
        documents_crees.append(doc_fr)
    except Exception as e:
        print(f"❌ Erreur lors de la génération du document français: {e}")

    print("\n" + "=" * 80)

    # Générer le document anglais
    try:
        doc_en = create_word_document(theme_name, "anglais")
        documents_crees.append(doc_en)
    except Exception as e:
        print(f"❌ Erreur lors de la génération du document anglais: {e}")

    return documents_crees

def main():
    """Point d'entrée principal"""
    parser = argparse.ArgumentParser(description='Génère un document Word pour un thème donné')
    parser.add_argument('theme', help='Nom du thème (ex: corps_humain, meteo)')
    parser.add_argument('--preview', action='store_true', help='Affiche la configuration sans générer')
    parser.add_argument('--bilingue', action='store_true', help='Génère les documents en français et en anglais')
    parser.add_argument('--langue', choices=['francais', 'anglais'], default='francais',
                       help='Langue du document (par défaut: francais)')

    args = parser.parse_args()

    try:
        if args.preview:
            config = load_selection_config(args.theme)
            print(f"📋 Configuration du thème '{args.theme}':")
            print(f"   Titre: {config['titre']}")
            print(f"   Éléments: {len(config['elements'])}")
            print(f"   Colonnes: {config.get('colonnes', 3)}")
            if args.theme in THEME_TRANSLATIONS:
                print(f"   Traductions disponibles:")
                for langue, titre in THEME_TRANSLATIONS[args.theme].items():
                    print(f"     {langue}: {titre}")
            return

        if args.bilingue:
            # Générer les deux versions
            documents_crees = create_bilingual_documents(args.theme)
            print(f"\n💡 {len(documents_crees)} document(s) créé(s):")
            for doc in documents_crees:
                print(f"   📄 {doc}")
        else:
            # Générer une seule version
            output_path = create_word_document(args.theme, args.langue)
            print(f"\n💡 Le document est prêt : {output_path}")

    except FileNotFoundError as e:
        print(f"❌ Erreur: {e}")
        print(f"💡 Téléchargez d'abord les images avec:")
        print(f"   python telecharger_images.py {args.theme}")
    except Exception as e:
        print(f"❌ Erreur: {e}")

if __name__ == "__main__":
    main()
