#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de documents Word thématiques en macédonien

Usage:
    python generer_document_theme.py corps_humain
    python generer_document_theme.py nourriture

Structure des dossiers attendue:
    themes/{theme_name}/
    ├── photos/              (toutes les images téléchargées)
    ├── selection.json       (config du thème)
    └── final/              (document généré)
"""

import sys
import json
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def charger_selection(theme_dir):
    """Charge le fichier selection.json du thème"""
    selection_path = theme_dir / "selection.json"
    with open(selection_path, 'r', encoding='utf-8') as f:
        return json.load(f)

def generer_document(theme_name):
    """Génère le document Word pour un thème donné"""

    # Chemins
    base_dir = Path(__file__).parent
    theme_dir = base_dir / "themes" / theme_name
    photos_dir = theme_dir / "photos"

    # Vérifications
    if not theme_dir.exists():
        print(f"❌ Erreur : Le thème '{theme_name}' n'existe pas")
        print(f"   Chemin attendu : {theme_dir}")
        return False

    if not photos_dir.exists():
        print(f"❌ Erreur : Le dossier photos n'existe pas : {photos_dir}")
        return False

    # Charger la configuration
    try:
        config = charger_selection(theme_dir)
    except FileNotFoundError:
        print(f"❌ Erreur : Fichier selection.json introuvable dans {theme_dir}")
        return False
    except json.JSONDecodeError as e:
        print(f"❌ Erreur JSON dans selection.json : {e}")
        return False

    # Créer le document
    doc = Document()

    # Titre (optionnel)
    titre = config.get('titre', f'Document {theme_name}')
    # title = doc.add_heading(titre, 0)
    # title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Calculer le nombre de lignes nécessaires
    elements = config['elements']
    colonnes = config.get('colonnes', 3)
    lignes = (len(elements) + colonnes - 1) // colonnes  # arrondi supérieur

    # Créer le tableau
    table = doc.add_table(rows=lignes, cols=colonnes)
    table.style = 'Table Grid'

    # Remplir le tableau
    idx = 0
    for row in table.rows:
        for cell in row.cells:
            if idx < len(elements):
                element = elements[idx]
                image_file = element['image_selectionnee']
                nom_macedonien = element['nom_macedonien']
                image_path = photos_dir / image_file

                # Vérifier que l'image existe
                if not image_path.exists():
                    print(f"⚠️  Attention : Image introuvable : {image_path}")
                    idx += 1
                    continue

                # Vider la cellule
                cell.text = ''

                # Ajouter l'image
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(2.0))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Ajouter un saut de ligne
                cell.add_paragraph()

                # Ajouter le nom en macédonien
                text_para = cell.add_paragraph(nom_macedonien)
                text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Formatter le texte
                for run in text_para.runs:
                    run.font.size = Pt(14)
                    run.font.bold = False
                    run.font.color.rgb = RGBColor(0, 0, 0)

                idx += 1

    # Nom du fichier de sortie
    output_filename = f"{theme_name.replace('_', ' ').title()}.docx"
    output_path = theme_dir / output_filename

    # Sauvegarder
    doc.save(str(output_path))

    # Rapport
    print(f"✅ Document créé avec succès !")
    print(f"📄 Titre : {titre}")
    print(f"📊 {len(elements)} éléments")
    print(f"📐 Tableau {lignes} × {colonnes}")
    print(f"📁 Emplacement : {output_path.absolute()}")

    return True

def main():
    """Point d'entrée du script"""
    if len(sys.argv) < 2:
        print("Usage : python generer_document_theme.py <nom_du_theme>")
        print("\nExemples :")
        print("  python generer_document_theme.py corps_humain")
        print("  python generer_document_theme.py nourriture")
        print("\nThèmes disponibles :")
        themes_dir = Path(__file__).parent / "themes"
        if themes_dir.exists():
            for theme in themes_dir.iterdir():
                if theme.is_dir():
                    print(f"  - {theme.name}")
        else:
            print("  (aucun thème trouvé)")
        sys.exit(1)

    theme_name = sys.argv[1]
    success = generer_document(theme_name)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()

